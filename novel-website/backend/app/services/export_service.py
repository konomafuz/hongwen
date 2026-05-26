import os
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from app.config import settings


EXPORT_DIR = os.path.join(os.path.dirname(os.path.dirname(os.path.dirname(__file__))), "exports")
os.makedirs(EXPORT_DIR, exist_ok=True)


async def export_to_word(project, chapters, include_outline=True) -> str:
    """Export project chapters to a Word document."""
    doc = Document()

    # Title page
    title_para = doc.add_heading(project.title, level=0)
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Project info
    info_para = doc.add_paragraph()
    info_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    info_run = info_para.add_run(
        f"字数：{project.word_count}  |  状态：{'进行中' if project.status == 'drafting' else '已完成'}"
    )
    info_run.font.size = Pt(11)
    info_run.font.color.rgb = RGBColor(128, 128, 128)

    doc.add_page_break()

    # Table of Contents
    doc.add_heading("目录", level=1)
    for ch in chapters:
        status_icon = "✓" if ch.status == "completed" else "○"
        doc.add_paragraph(f"第{ch.chapter_number}章  {ch.title}  ({status_icon})")

    doc.add_page_break()

    # Chapters
    for i, ch in enumerate(chapters, 1):
        doc.add_heading(f"第{ch.chapter_number}章  {ch.title}", level=1)

        # Include outline if requested
        if include_outline and ch.outline:
            outline_para = doc.add_paragraph()
            outline_run = outline_para.add_run("【本章大纲】")
            outline_run.bold = True
            outline_run.font.color.rgb = RGBColor(100, 100, 100)
            doc.add_paragraph(ch.outline)
            doc.add_paragraph("---")

        # Content
        if ch.content:
            doc.add_paragraph(ch.content)

        # Page break between chapters
        if i < len(chapters):
            doc.add_page_break()

    # Save
    safe_title = "".join(c for c in project.title if c.isalnum() or c in " _-").strip()
    filename = f"{safe_title}_第1-{len(chapters)}章.docx"
    if not filename.endswith(".docx"):
        filename += ".docx"

    filepath = os.path.join(EXPORT_DIR, filename)
    doc.save(filepath)
    return filepath