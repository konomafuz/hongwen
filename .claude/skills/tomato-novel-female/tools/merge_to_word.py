# -*- coding: utf-8 -*-
"""
自动合并章节 Markdown 文件并生成排版精美的 Word 文档。
作者：Antigravity
"""

import os
import re
import sys
import glob
import subprocess
import json

# ----------------- 依赖检测与安装 -----------------
try:
    import docx
except ImportError:
    print("未找到 python-docx 库，正在自动安装...")
    try:
        subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx"])
        import docx
        print("python-docx 库安装成功！\n")
    except Exception as e:
        print(f"安装 python-docx 失败，请手动运行 'pip install python-docx' 后重试。错误信息: {e}")
        sys.exit(1)

from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

# ----------------- 工具函数 -----------------
def get_workspace_dir():
    """动态获取包含 novel_core_settings.md 的工作区根目录"""
    if os.path.exists(os.path.join(os.getcwd(), "novel_core_settings.md")):
        return os.getcwd()
    current_dir = os.path.dirname(os.path.abspath(__file__))
    workspace_dir = current_dir
    while workspace_dir:
        if os.path.exists(os.path.join(workspace_dir, "novel_core_settings.md")):
            return workspace_dir
        parent = os.path.dirname(workspace_dir)
        if parent == workspace_dir:
            break
        workspace_dir = parent
    return os.getcwd()  # 兜底返回当前运行目录

def get_novel_title(workspace_dir):
    """从核心设定文件中提取小说标题"""
    default_title = "绑定吃瓜系统，我靠剧透爆红了"
    settings_path = os.path.join(workspace_dir, "novel_core_settings.md")
    if os.path.exists(settings_path):
        try:
            with open(settings_path, "r", encoding="utf-8") as f:
                first_line = f.readline().strip()
                # 提取《书名》
                match = re.search(r"《([^》]+)》", first_line)
                if match:
                    return match.group(1)
        except Exception as e:
            print(f"提示: 读取小说标题失败，将使用默认书名。错误: {e}")
    return default_title

def set_font_yahei(run):
    """为文档中的 Run 元素强制应用微软雅黑字体（中英文一致）"""
    run.font.name = "Microsoft YaHei"
    rPr = run.element.get_or_add_rPr()
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:ascii'), 'Microsoft YaHei')
    rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
    rFonts.set(qn('w:hAnsi'), 'Microsoft YaHei')
    rPr.append(rFonts)

def add_chapter_title(doc, title_text, is_first=False):
    """添加规范化的章节标题并居中"""
    if not is_first:
        doc.add_page_break()  # 每章开头自动分页
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(36)  # 标题上方间距
    p.paragraph_format.space_after = Pt(24)   # 标题下方间距
    p.paragraph_format.keep_with_next = True   # 防止标题和正文分离（孤行控制）
    
    run = p.add_run(title_text)
    run.bold = True
    run.font.size = Pt(16)  # 三号字 (16pt)
    set_font_yahei(run)

def add_body_paragraph(doc, text, book_type="female_novel"):
    """添加正文段落，并根据作品体裁进行智能排版"""
    text_clean = text.strip()
    if not text_clean:
        return
    
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.25         # 1.25倍行距
    p.paragraph_format.space_after = Pt(6)         # 段后间距 6pt
    p.paragraph_format.space_before = Pt(0)
    
    if book_type == "short_drama":
        p.paragraph_format.first_line_indent = Pt(0)  # 剧本无首行缩进
        
        # 1. 幕标识: **第1幕：【场景】大厅 ...**
        if text_clean.startswith('**') and text_clean.endswith('**'):
            clean_text = text_clean[2:-2]
            run = p.add_run(clean_text)
            run.bold = True
            run.font.size = Pt(13)  # 幕标识稍微大一点 (小四)
            set_font_yahei(run)
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after = Pt(6)
            return
            
        # 2. 动作/神态指示: （苏绿茶脸色一白...） 或 (苏绿茶脸色一白...)
        elif (text_clean.startswith('（') and text_clean.endswith('）')) or (text_clean.startswith('(') and text_clean.endswith(')')):
            run = p.add_run(text_clean)
            run.italic = True
            # 设置颜色为灰色 (RGB: 128, 128, 128)
            run.font.color.rgb = docx.shared.RGBColor(128, 128, 128)
            set_font_yahei(run)
            p.paragraph_format.left_indent = Inches(0.4)  # 缩进表示动作
            return
            
        # 3. 台词: 角色名：台词内容
        match = re.match(r"^([^：\s]+[：:])(.*)$", text_clean)
        if match:
            char_part = match.group(1)
            dialogue_part = match.group(2)
            
            run_char = p.add_run(char_part)
            run_char.bold = True
            set_font_yahei(run_char)
            
            # 对台词部分可能存在的括号神态指示（如：(哭腔)安然...）进行处理
            sub_parts = re.split(r'([（\(].*?[）\)])', dialogue_part)
            for sub_part in sub_parts:
                if not sub_part:
                    continue
                if (sub_part.startswith('（') and sub_part.endswith('）')) or (sub_part.startswith('(') and sub_part.endswith(')')):
                    run_sub = p.add_run(sub_part)
                    run_sub.italic = True
                    run_sub.font.color.rgb = docx.shared.RGBColor(128, 128, 128)
                    set_font_yahei(run_sub)
                else:
                    run_sub = p.add_run(sub_part)
                    set_font_yahei(run_sub)
            return

    # 默认/长篇小说排版
    p.paragraph_format.first_line_indent = Pt(24)  # 12pt字号下，2字符首行缩进 = 24pt
    # 简单 Markdown 格式解析：粗体 (**text**) 与 斜体 (*text*)
    parts = re.split(r'(\*\*.*?\*\*|\*.*?\*)', text_clean)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run = p.add_run(part[2:-2])
            run.bold = True
        elif part.startswith('*') and part.endswith('*'):
            run = p.add_run(part[1:-1])
            run.italic = True
        else:
            run = p.add_run(part)
        set_font_yahei(run)

# ----------------- 主程序入口 -----------------
def main():
    workspace_dir = get_workspace_dir()
    print(f"定位到工作区根目录: {workspace_dir}")
    
    # 加载项目状态以获取作品类型
    book_type = "female_novel"
    state_path = os.path.join(workspace_dir, "novel_project.json")
    if os.path.exists(state_path):
        try:
            with open(state_path, "r", encoding="utf-8") as sf:
                state_data = json.load(sf)
                book_type = state_data.get("book_type", "female_novel")
            print(f"检测到作品品类: {book_type}")
        except Exception as e:
            print(f"提示: 读取 novel_project.json 失败，默认为 female_novel。错误: {e}")
            
    # 获取小说标题
    novel_title = get_novel_title(workspace_dir)
    print(f"检测到小说标题: 《{novel_title}》")
    
    # 搜索并对章节文件进行数值排序
    chapter_pattern = os.path.join(workspace_dir, "chapter_*.md")
    chapter_files = glob.glob(chapter_pattern)
    
    if not chapter_files:
        print(f"错误: 未在目录中找到任何符合 'chapter_*.md' 的文件！")
        return
    
    # 提取数字排序，例如 chapter_01.md, chapter_2.md, chapter_10.md 等
    def get_chap_num(filepath):
        basename = os.path.basename(filepath)
        match = re.search(r"chapter_(\d+)", basename)
        return int(match.group(1)) if match else 0
    
    chapter_files.sort(key=get_chap_num)
    print(f"找到以下 {len(chapter_files)} 个章节文件，正在准备按顺序合并:")
    for f in chapter_files:
        print(f" - {os.path.basename(f)}")
    
    # 初始化 Word 文档
    doc = docx.Document()
    
    # 设置页面边距（上下左右 1 英寸，即 2.54 厘米）
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
    
    # 设置默认段落样式字体为微软雅黑
    style = doc.styles['Normal']
    style.font.name = 'Microsoft YaHei'
    style.font.size = Pt(12)
    
    # 循环读取合并
    for idx, filepath in enumerate(chapter_files):
        print(f"正在读取并导入: {os.path.basename(filepath)}...")
        try:
            with open(filepath, "r", encoding="utf-8") as f:
                lines = f.readlines()
        except Exception as e:
            print(f"读取章节文件失败: {filepath}. 错误: {e}")
            continue
        
        if not lines:
            continue
        
        # 提取第一行作为章节标题并规范化
        title_line = lines[0].strip()
        
        # 正则解析 "# 第X章/集 章节名" 或类似格式
        match = re.match(r"^#\s*第\s*(\d+)\s*[章回集]\s*[:：\s]*\s*(.*)$", title_line)
        if match:
            chap_num = match.group(1)
            chap_name = match.group(2).strip()
            unit_name = "集" if book_type == "short_drama" else "章"
            normalized_title = f"第{chap_num}{unit_name}：{chap_name}"
        else:
            # 兜底清理 # 符号
            normalized_title = title_line.lstrip("#").strip()
        
        # 写入标题到Word
        add_chapter_title(doc, normalized_title, is_first=(idx == 0))
        
        # 写入正文
        in_body = False
        for line in lines[1:]:
            # 跳过空行或仅含分隔符的行
            line_str = line.strip()
            if not line_str or line_str in ["---", "***", "——"]:
                continue
            add_body_paragraph(doc, line_str, book_type=book_type)
            
    # 输出文件路径
    output_filename = f"{novel_title}.docx"
    output_path = os.path.join(workspace_dir, output_filename)
    
    try:
        doc.save(output_path)
        print(f"\n合并成功！Word 文档已保存至:\n{output_path}")
    except Exception as e:
        print(f"保存 Word 文档失败。错误信息: {e}")

if __name__ == "__main__":
    main()
